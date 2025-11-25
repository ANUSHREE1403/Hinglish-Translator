import argparse
import os
import time
import logging
from typing import Optional, List, Tuple

# External libraries
from dotenv import load_dotenv  # type: ignore

# Globals for engine selection
ENGINE = "gemini"
ARGOS_TRANSLATOR = None


def configure_gemini(api_key: str, model_name: str):
    # Lazy import so Argos-only runs don't require Google libs configured
    import google.generativeai as genai  # type: ignore
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(model_name)


def ensure_argos_translator(source_lang_code: str, target_lang_code: str):
    global ARGOS_TRANSLATOR
    try:
        import argostranslate.package as argos_pkg  # type: ignore
        import argostranslate.translate as argos_tr  # type: ignore
    except Exception as e:  # noqa: BLE001
        raise SystemExit(
            "Argos Translate not installed. Please run: pip install argostranslate"
        ) from e

    # If already installed and loaded, reuse
    if ARGOS_TRANSLATOR is not None:
        return ARGOS_TRANSLATOR

    # Make sure the language package is installed
    try:
        installed = argos_pkg.get_installed_packages()
        if not any(p.from_code == source_lang_code and p.to_code == target_lang_code for p in installed):
            logging.info(
                "Argos package %s->%s not found; downloading and installing (one-time).",
                source_lang_code,
                target_lang_code,
            )
            available = argos_pkg.get_available_packages()
            pkg = next(
                p for p in available if p.from_code == source_lang_code and p.to_code == target_lang_code
            )
            pkg_path = pkg.download()
            argos_pkg.install_from_path(pkg_path)
    except StopIteration as e:  # noqa: BLE001
        raise SystemExit(
            f"Argos language pair {source_lang_code}->{target_lang_code} not available."
        ) from e

    ARGOS_TRANSLATOR = argos_tr
    return ARGOS_TRANSLATOR


def normalize_translated_text(text: str) -> str:
    """Normalize LLM output to a single clean line suitable for the Translation cell."""
    if not text:
        return ""
    # Collapse newlines and excessive whitespace
    cleaned = " ".join(text.replace("\n", " ").split())
    # Strip wrapping quotes or stray punctuation
    if len(cleaned) >= 2 and ((cleaned[0] == '"' and cleaned[-1] == '"') or (cleaned[0] == "'" and cleaned[-1] == "'")):
        cleaned = cleaned[1:-1].strip()
    # Remove markdown asterisks and common speaker prefixes
    cleaned = cleaned.replace("**", "").strip()
    for prefix in ["JJ:", "Mikey:", "JJ -", "Mikey -", "JJ —", "Mikey —"]:
        if cleaned.lower().startswith(prefix.lower()):
            cleaned = cleaned[len(prefix):].strip()
    return cleaned


def contains_devanagari(text: str) -> bool:
    return any('\u0900' <= ch <= '\u097F' for ch in text)


def translate_text(model, text: str, target_language: str, tone: str = "natural", style_prompt: str = "") -> str:
    if not text or not text.strip():
        return ""

    base_prompt = (
        f"Task: Translate the following line into Hinglish (Roman Hindi in English letters only).\n"
        f"Constraints: kid-friendly Maizen tone, energetic; use ONLY A-Z letters, digits, and basic punctuation; no Devanagari or emojis;\n"
        f"male pronouns for JJ and Mikey; similar reading time as source; NO speaker names;\n"
        f"Output: return exactly ONE single-line Hinglish sentence, nothing else."
    )
    system_prompt = (style_prompt.strip() or base_prompt)

    if ENGINE == "argos":
        # Argos path (offline)
        # Map target language name to code; default to 'hi' for Hindi
        target_code = "hi" if target_language.strip().lower().startswith("hin") else target_language.strip().lower()
        src_code = "en"
        translator = ensure_argos_translator(src_code, target_code)
        try:
            result = translator.translate(text.strip(), src_code, target_code)
        except TypeError:
            # Older argostranslate uses translate(text, from_code, to_code) or sets languages globally
            result = translator.translate(text.strip(), src_code, target_code)
        return (result or "").strip()
    else:
        # Gemini path (online)
        # Retry with simple exponential backoff on transient errors
        delays = [0.5, 1, 2, 4]
        for attempt, delay in enumerate([0] + delays):
            if delay:
                logging.debug(f"Retry backoff sleeping {delay:.1f}s before attempt {attempt+1}")
                time.sleep(delay)
            try:
                prompt = f"{system_prompt}\nSource: \"{text.strip()}\"\nHinglish:"
                logging.debug("Calling Gemini (attempt %d).", attempt + 1)
                response = model.generate_content(prompt)
                candidate = (response.text or "").strip()
                candidate = normalize_translated_text(candidate)
                if contains_devanagari(candidate):
                    logging.debug("Candidate contained Devanagari; requesting strict roman-only rewrite.")
                    strict = (
                        "Rewrite strictly in Hinglish (English letters only). No Hindi script, no speaker names, one single line only.\n"
                        f"Source: \"{text.strip()}\"\nHinglish:"
                    )
                    response2 = model.generate_content(strict)
                    candidate2 = normalize_translated_text((response2.text or "").strip())
                    if not contains_devanagari(candidate2):
                        candidate = candidate2
                logging.debug(
                    "Received translation (attempt %d). Output preview: %s",
                    attempt + 1,
                    (candidate[:120] + "...") if len(candidate) > 120 else candidate,
                )
                return candidate
            except Exception as e:  # noqa: BLE001
                last_error = e
                logging.warning("Gemini call failed on attempt %d: %s", attempt + 1, str(e))
                continue
        raise RuntimeError(f"Gemini translation failed after retries: {last_error}")


def try_process_docx(
    input_path: str,
    output_path: str,
    source_column: str,
    target_column: str,
    model,
    target_language: str,
    tone: str,
    delay_seconds: float,
    max_rows: int,
    style_prompt: str,
) -> bool:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return False

    if not input_path.lower().endswith(".docx"):
        return False

    doc = Document(input_path)
    modified = False

    table_index = -1
    total_rows = 0
    total_translated = 0
    total_skipped = 0
    for table in doc.tables:
        table_index += 1
        # Find header row if present by matching columns
        header_idx = 0
        if table.rows:
            header_cells = table.rows[0].cells
            header_map = {c.text.strip().lower(): i for i, c in enumerate(header_cells)}
            src_idx = header_map.get(source_column.strip().lower())
            tgt_idx = header_map.get(target_column.strip().lower())
            logging.info(
                "DOCX table %d: header map=%s, src_idx=%s, tgt_idx=%s",
                table_index,
                header_map,
                src_idx,
                tgt_idx,
            )

            # If headers not matched, fallback to positional: assume first is source, second is target
            if src_idx is None or tgt_idx is None:
                if len(header_cells) >= 2:
                    src_idx, tgt_idx = 0, 1
                else:
                    continue

            for row_i, row in enumerate(table.rows[1:], start=2):
                cells = row.cells
                if max(src_idx, tgt_idx) >= len(cells):
                    continue
                source_text = cells[src_idx].text
                current_target = cells[tgt_idx].text
                if source_text and not current_target.strip():
                    logging.info(
                        "DOCX table %d row %d: translating (delay %.1fs)",
                        table_index,
                        row_i,
                        delay_seconds,
                    )
                    if delay_seconds > 0:
                        time.sleep(delay_seconds)
                    translated = translate_text(model, source_text, target_language, tone, style_prompt)
                    cells[tgt_idx].text = normalize_translated_text(translated)
                    modified = True
                    total_translated += 1
                    if max_rows > 0 and total_translated >= max_rows:
                        logging.info("DOCX max_rows reached: %d", max_rows)
                        break
                else:
                    total_skipped += 1
                total_rows += 1
            if max_rows > 0 and total_translated >= max_rows:
                break

    if modified:
        doc.save(output_path)
    else:
        # If nothing changed, still save a copy to keep behavior consistent
        doc.save(output_path)
    logging.info(
        "DOCX summary: tables=%d rows=%d translated=%d skipped=%d output=%s",
        len(doc.tables),
        total_rows,
        total_translated,
        total_skipped,
        output_path,
    )
    return True


def try_process_xlsx(
    input_path: str,
    output_path: str,
    source_column: str,
    target_column: str,
    model,
    target_language: str,
    tone: str,
    delay_seconds: float,
    max_rows: int,
    style_prompt: str,
) -> bool:
    try:
        import openpyxl  # type: ignore
    except Exception:
        return False

    if not input_path.lower().endswith((".xlsx", ".xlsm")):
        return False

    wb = openpyxl.load_workbook(input_path)
    for ws in wb.worksheets:
        # Header row assumed at row 1
        headers = {}
        for col_idx, cell in enumerate(ws[1], start=1):
            headers[str(cell.value).strip().lower() if cell.value else ""] = col_idx

        src_idx = headers.get(source_column.strip().lower())
        tgt_idx = headers.get(target_column.strip().lower())
        if src_idx is None or tgt_idx is None:
            # Skip sheet if headers not found
            continue

        translated_count = 0
        skipped_count = 0
        for row_idx in range(2, ws.max_row + 1):
            source_value = ws.cell(row=row_idx, column=src_idx).value
            target_value = ws.cell(row=row_idx, column=tgt_idx).value
            if source_value and (not str(target_value).strip()):
                logging.info(
                    "XLSX sheet '%s' row %d: translating (delay %.1fs)",
                    ws.title,
                    row_idx,
                    delay_seconds,
                )
                if delay_seconds > 0:
                    time.sleep(delay_seconds)
                translated = translate_text(model, str(source_value), target_language, tone, style_prompt)
                ws.cell(row=row_idx, column=tgt_idx, value=normalize_translated_text(translated))
                translated_count += 1
                if max_rows > 0 and translated_count >= max_rows:
                    logging.info("XLSX max_rows reached: %d", max_rows)
                    break
            else:
                skipped_count += 1
        logging.info(
            "XLSX sheet '%s' summary: rows=%d translated=%d skipped=%d",
            ws.title,
            ws.max_row - 1,
            translated_count,
            skipped_count,
        )
        if max_rows > 0 and translated_count >= max_rows:
            break

    wb.save(output_path)
    return True


def try_process_csv(
    input_path: str,
    output_path: str,
    source_column: str,
    target_column: str,
    model,
    target_language: str,
    tone: str,
    delay_seconds: float,
    max_rows: int,
    style_prompt: str,
) -> bool:
    try:
        import pandas as pd  # type: ignore
    except Exception:
        return False

    if not input_path.lower().endswith(".csv"):
        return False

    df = pd.read_csv(input_path)
    if source_column not in df.columns or target_column not in df.columns:
        return False

    translated_count = 0
    skipped_count = 0

    def fill_row(val, existing, idx):
        if pd.isna(val):
            return existing
        if (existing is None) or (isinstance(existing, float) and pd.isna(existing)) or (str(existing).strip() == ""):
            logging.info("CSV row %d: translating (delay %.1fs)", idx + 2, delay_seconds)
            if delay_seconds > 0:
                time.sleep(delay_seconds)
            result = translate_text(model, str(val), target_language, tone, style_prompt)
            result = normalize_translated_text(result)
            nonlocal translated_count
            translated_count += 1
            return result
        return existing

    out_values = []
    for i in range(len(df)):
        if max_rows > 0 and translated_count >= max_rows:
            out_values.append(df[target_column].iloc[i])
            skipped_count += 1
            continue
        new_val = fill_row(df[source_column].iloc[i], df[target_column].iloc[i], i)
        if new_val == df[target_column].iloc[i]:
            skipped_count += 1
        out_values.append(new_val)
    df[target_column] = out_values
    df.to_csv(output_path, index=False)
    logging.info(
        "CSV summary: rows=%d translated=%d skipped=%d output=%s",
        len(df),
        translated_count,
        skipped_count,
        output_path,
    )
    return True


def main():
    parser = argparse.ArgumentParser(description="Translate dialogue column to target language and fill translation column using Gemini.")
    parser.add_argument("input", help="Input file (.docx, .xlsx, .csv)")
    parser.add_argument("--output", help="Output file path (defaults to <input>.translated<ext>)")
    parser.add_argument("--api-key", dest="api_key", help="Gemini API key (or set GOOGLE_API_KEY env)")
    parser.add_argument("--engine", choices=["gemini", "argos"], default="argos", help="Translation engine: online Gemini or offline Argos")
    parser.add_argument("--model", default="gemini-flash-latest", help="Gemini model name (engine=gemini)")
    parser.add_argument("--env-file", default="", help="Path to .env file to load (defaults to .env if present)")
    parser.add_argument("--api-key-env-var", default="GOOGLE_API_KEY", help="Env var name to read Gemini API key from")
    parser.add_argument("--source-column", default="Dialogue", help="Name of the source/dialogue column")
    parser.add_argument("--target-column", default="Translation", help="Name of the translation column to fill")
    parser.add_argument("--target-language", default="Hindi", help="Target language (e.g., Hindi)")
    parser.add_argument("--tone", default="natural", help="Desired tone (e.g., natural, formal, colloquial)")
    parser.add_argument("--delay-seconds", type=float, default=7.0, help="Delay between requests to respect rate limits")
    parser.add_argument("--max-rows", type=int, default=0, help="Translate only this many rows (per file). 0 = all")
    parser.add_argument(
        "--style-prompt",
        default=(
            """
Read through the Dialogue column and add really creative, engaging, and kid-friendly Hinglish translation/transcreation into the Translation column. Make sure the intent remains the same. Structure sentences so the reading time is similar to the source (dub‑friendly, natural flow). Use correct male pronouns for JJ and Mikey.

Here's some of the scripting I do for your reference; follow this style:

Alright. Let's get started. Today we're gonna clear every stage of the Twilight Forest. There are a lot of tough bosses we're gonna have to defeat.
Chalo, shuru karte hain! Toh aaj hum Twilight Forest ke har stage ko clear karenge. Bahut saare khatarnaak bosses ko harana padega.
Really? Are we gonna be able to beat them all?
Sach mein? Kya hum sabko hara paayenge?
We're gonna be fine, and that's because...
Hum sabki band baja denge, aur woh isliye kyunki...
(R) Hm?
(R) Hmm?
This! Inside this chest is the most powerful equipment in the game. And it's all ours. We'd better get ready. Armor up!
Yeh! Is chest ke andar game ka sabse powerful hathiyaar hai. Aur yeh sab hamara hai. Toh ready ho jao. Armor pehen lo!
Wow! With this, we'll be unstoppable.
Wah! Aise toh humein koi nahi rok payega
That's because this is the strongest armor. Take a look at these weapons, too.
Kyunki yeh sabse mazboot armor hai. Aur in hathiyaaro ko toh dekho
(R) Hoh! Powerful.
(R) Hoh! Khatarnaak bhai.
These weapons are super strong. I guess I should explain how our equipment works. Our armor gives us the ability to fly.
Yeh hathiyaar super strong hain. Chalo bata doon ki hamara equipment kaise kaam karta hai. Hamara armor hamein udne ki ability deta hai.
Wow, that's incredible.
Wah, kya bawaal cheez hai.
Plus it gives us super speed and a tough defense. That's it. Let's head into the Twilight Forest.
Plus yeh hamein super speed aur ek badhiya defense deta hai. Samajh gaye na? Toh chalo Twilight Forest chalte hain.
Let's go!
Chalo!
Go ahead and throw a diamond in there.
Ek diamond daal do usmein.
Sure. Is that how we get there? Three, two, one, go!
Sure. Aise jaate hain kya wahaan? Three, two, one, go!
Woah.
Woah.
Incredible! A portal. Let's go!
Incredible! Ek portal. Chalo!
And we're off!
Aur hum chal diye!
Nice! We've finally arrived in the Twilight Forest.
Nice! Hum finally Twilight Forest mein aa gaye.

For Context: We are localising the YouTuber Maizen; main characters are JJ and Mikey.

Output rules: Return ONLY the translated line; no extra lines/notes/speaker names. Translate into Hinglish (Roman Hindi using English letters only). Use English letters and basic punctuation only—no Devanagari, no emojis also please sure you try to keep the emglish words also in hindi like avoid words like creatures and use the hindi word for that but in english alphabets .
"""
        ),
        help="Custom style prompt for the translator (engine=gemini)",
    )
    parser.add_argument("--log-level", default="INFO", choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"], help="Logging level")

    args = parser.parse_args()

    logging.basicConfig(
        level=getattr(logging, args.log_level.upper(), logging.INFO),
        format="%(asctime)s | %(levelname)s | %(message)s",
        datefmt="%H:%M:%S",
    )

    global ENGINE
    ENGINE = args.engine

    api_key = None
    if ENGINE == "gemini":
        # If --api-key is provided, do NOT load any .env files.
        if args.api_key:
            api_key = args.api_key
        else:
            # Load .env automatically; prefer specified file if provided, but be resilient to encoding issues.
            if args.env_file:
                try:
                    load_dotenv(args.env_file, override=False)
                    logging.info("Loaded env file: %s", args.env_file)
                except Exception as e:  # noqa: BLE001
                    logging.warning("Skipping env file due to read error: %s", str(e))
            else:
                try:
                    load_dotenv(override=False)  # loads .env if it exists
                except Exception as e:  # noqa: BLE001
                    logging.warning("Skipping default .env due to read error: %s", str(e))
            api_key = os.getenv(args.api_key_env_var)
        if not api_key:
            raise SystemExit(
                f"Missing API key. Provide --api-key or set {args.api_key_env_var} (optionally via --env-file)."
            )

    input_path = args.input
    if not os.path.exists(input_path):
        raise SystemExit(f"Input not found: {input_path}")

    root, ext = os.path.splitext(input_path)
    output_path = args.output or f"{root}.translated{ext}"

    logging.info(
        "Starting translation | engine=%s model=%s target_language=%s tone=%s delay=%.1fs input=%s output=%s",
        args.engine,
        args.model,
        args.target_language,
        args.tone,
        args.delay_seconds,
        input_path,
        output_path,
    )
    model = None
    if ENGINE == "gemini":
        model = configure_gemini(api_key, args.model)
    else:
        # Ensure Argos translator exists and language pair is installed
        ensure_argos_translator("en", "hi" if args.target_language.strip().lower().startswith("hin") else args.target_language.strip().lower())

    # Try handlers in order based on extension
    handlers = [try_process_docx, try_process_xlsx, try_process_csv]
    for handler in handlers:
        ok = handler(
            input_path=input_path,
            output_path=output_path,
            source_column=args.source_column,
            target_column=args.target_column,
            model=model,
            target_language=args.target_language,
            tone=args.tone,
            delay_seconds=args.delay_seconds,
            max_rows=args.max_rows,
            style_prompt=args.style_prompt,
        )
        if ok:
            logging.info("Done. Wrote: %s", output_path)
            return

    raise SystemExit(
        "Unsupported format or missing libraries. Install extras and ensure columns exist. "
        "Supported: .docx (python-docx), .xlsx (openpyxl), .csv (pandas)."
    )


if __name__ == "__main__":
    main()


